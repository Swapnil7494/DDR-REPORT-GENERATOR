from docx import Document
import os

def create_docx(report_text):

    # Create output folder if it doesn't exist
    os.makedirs("output", exist_ok=True)

    doc = Document()

    doc.add_heading(
        "DDR Report",
        level=1
    )

    doc.add_paragraph(report_text)

    output_file = "output/DDR_Report.docx"

    doc.save(output_file)

    return output_file